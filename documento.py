import docx
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from src.models import Militar
from sqlalchemy import and_
from src import app  # Importe sua aplicação Flask

with app.app_context():
    # Adicionar prints para depuração
    coronel_qobm = Militar.query.filter(
        and_(
            Militar.posto_grad_id == 14,
            Militar.quadro_id == 2,
            Militar.especialidade_id == 3
        )
    ).count()
    print(f"Coronel QOBM count: {coronel_qobm}")

    tc_qobm = Militar.query.filter(
        and_(
            Militar.posto_grad_id == 13,
            Militar.quadro_id == 2,
            Militar.especialidade_id == 3
        )
    ).count()
    print(f"Tenente Coronel QOBM count: {tc_qobm}")

    major_qobm = Militar.query.filter(
        and_(
            Militar.posto_grad_id == 12,
            Militar.quadro_id == 2,
            Militar.especialidade_id == 3
        )
    ).count()
    print(f"Major QOBM count: {major_qobm}")

    capitao_qobm = Militar.query.filter(
        and_(
            Militar.posto_grad_id == 11,
            Militar.quadro_id == 2,
            Militar.especialidade_id == 3
        )
    ).count()
    print(f"Capitão QOBM count: {capitao_qobm}")

    ten_1 = Militar.query.filter(
        and_(
            Militar.posto_grad_id == 10,
            Militar.quadro_id == 2,
            Militar.especialidade_id == 3
        )
    ).count()
    print(f"1º Tenente QOBM count: {ten_1}")

    ten_2 = Militar.query.filter(
        and_(
            Militar.posto_grad_id == 9,
            Militar.quadro_id == 2,
            Militar.especialidade_id == 3
        )
    ).count()
    print(f"2º Tenente QOBM count: {ten_2}")

    # Caminho do seu template
    template_path = 'EFETIVO POR QUADRO_ESPECIALIDADE.docx'
    # Caminho do arquivo que será criado
    output_path = 'EFETIVO_POR_QUADRO_ESPECIALIDADE.docx'

    # Carregar o documento template
    doc = docx.Document(template_path)

    # Adicionar uma tabela
    # 3 colunas e 9 linhas para acomodar o título, subtítulo e as novas informações
    table = doc.add_table(rows=9, cols=3)
    table.style = 'Table Grid'

    # Adicionar o título
    cell = table.cell(0, 0)
    cell.merge(table.cell(0, 2))  # Mesclar células da primeira linha
    cell.text = 'DIRETORIA DE RECURSOS HUMANOS/CBMAM'
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(16)
    run.bold = True

    # Adicionar o subtítulo
    cell = table.cell(1, 0)
    cell.merge(table.cell(1, 1))  # Mesclar as duas primeiras células da segunda linha
    cell.text = 'EFETIVO DOS COMBATENTES/RESUMO'
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(14)
    run.bold = True

    # Adicionar "EXISTENTE" na célula ao lado do subtítulo
    cell = table.cell(1, 2)
    cell.text = 'EXISTENTE'
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(14)
    run.bold = True

    # Adicionar os postos e quadro na tabela
    postos = ["CORONEL", "TENENTE CORONEL", "MAJOR", "CAPITÃO", "1º TENENTE", "2º TENENTE"]

    # Mesclar células na coluna B de A3 até A8
    cell = table.cell(2, 1)
    cell.merge(table.cell(7, 1))
    cell.text = "QOBM"
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(14)

    # Adicionar postos na coluna A
    for i, posto in enumerate(postos, start=2):
        table.cell(i, 0).text = posto
        paragraph = table.cell(i, 0).paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = paragraph.runs[0]
        run.font.size = Pt(14)

    # Adicionar o resultado do filtro na célula correspondente
    # Coronel QOBM
    cell = table.cell(2, 2)
    cell.text = str(coronel_qobm)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(14)

    # Tenente Coronel QOBM
    cell = table.cell(3, 2)
    cell.text = str(tc_qobm)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(14)

    # Major QOBM
    cell = table.cell(4, 2)
    cell.text = str(major_qobm)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(14)

    # Capitão QOBM
    cell = table.cell(5, 2)
    cell.text = str(capitao_qobm)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(14)

    # 1º Tenente QOBM
    cell = table.cell(6, 2)
    cell.text = str(ten_1)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(14)

    # 2º Tenente QOBM
    cell = table.cell(7, 2)
    cell.text = str(ten_2)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0]
    run.font.size = Pt(14)

    # Salvar o documento modificado
    doc.save(output_path)
    print(f"Documento '{output_path}' criado com sucesso!")
